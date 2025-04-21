#%%
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
from io import BytesIO
import pandas as pd

# class Document():
#     def __init__(self):
#         self = Doc()

def add_dataframe(doc: Document, df: pd.DataFrame, transpose=False):
    if transpose:
        t = doc.add_table(df.shape[1] + 1, df.shape[0] + 1)
        # Add the column headers
        for j in range(df.shape[0]):
            t.cell(0, j + 1).text = str(df.index[j])
        t.cell(0, 0).text = ""
        # Add the row headers
        for i in range(df.shape[1]):
            t.cell(i + 1, 0).text = str(df.columns[i])
        # Add the DataFrame values
        for i in range(df.shape[1]):
            for j in range(df.shape[0]):
                t.cell(i + 1, j + 1).text = str(df.values[j, i])
    else:
        t = doc.add_table(df.shape[0] + 1, df.shape[1])
        # Add the column headers
        for j in range(df.shape[1]):
            t.cell(0, j).text = str(df.columns[j])
        # Add the DataFrame values
        for i in range(df.shape[0]):
            for j in range(df.shape[1]):
                t.cell(i + 1, j).text = str(df.values[i, j])

    return t

def add_figure(doc: Document, fig, ax, caption=None, add_break=True):
    memfile = BytesIO()
    # fig.tight_layout()
    fig.savefig(memfile)
    if caption: 
        doc.add_paragraph(caption)
    doc.add_picture(memfile, width=Inches(6))
    plt.close(fig)
    
    if add_break:
        doc.add_page_break()

# document = Document()
#%%